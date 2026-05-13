#!/usr/bin/env python3
"""Post-process pandoc-generated proposal: DM Sans, black headings, h2 accent bar."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOCX_PATH = "/Users/tarunrastogi/Tarun Working Env/merrakii/docs/proposal-cosmetic-brand-website.docx"
# Gradient bar (same as h2::before in HTML); bundled next to this script
BAR_IMAGE = str(Path(__file__).resolve().parent / "proposal-h2-accent.png")

FONT = "DM Sans"
BLACK = RGBColor(0x00, 0x00, 0x00)
INK = RGBColor(0x1A, 0x1D, 0x24)


def set_rfonts_on_run(run, font_name: str) -> None:
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def apply_font_all_runs(doc: Document, font_name: str) -> None:
    def proc_paragraph(p):
        for r in p.runs:
            set_rfonts_on_run(r, font_name)
            r.font.name = font_name

    for p in doc.paragraphs:
        proc_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    proc_paragraph(p)


def clear_paragraph_runs(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            p.remove(child)


def add_heading2_with_bar(paragraph, bar_path: str) -> None:
    text = paragraph.text
    style = paragraph.style
    clear_paragraph_runs(paragraph)

    run_img = paragraph.add_run()
    run_img.add_picture(bar_path, width=Pt(4.5), height=Pt(16))
    set_rfonts_on_run(run_img, FONT)

    run_txt = paragraph.add_run(" " + text)
    run_txt.bold = True
    run_txt.font.color.rgb = BLACK
    set_rfonts_on_run(run_txt, FONT)

    paragraph.style = style


def style_doc_defaults(doc: Document) -> None:
    mapping = [
        ("Normal", INK, None),
        ("Body Text", INK, None),
        ("First Paragraph", INK, None),
        ("Title", BLACK, Pt(22)),
        ("Heading 1", BLACK, Pt(18)),
        ("Heading 2", BLACK, Pt(13)),
        ("Heading 3", BLACK, Pt(12)),
        ("Heading 4", BLACK, Pt(11)),
        ("List Paragraph", INK, None),
    ]
    for name, color, size in mapping:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = FONT
        st.font.color.rgb = color
        if size is not None:
            st.font.size = size

        # Theme fonts: set on style element if present
        rpr = st.element.rPr
        if rpr is not None:
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:ascii"), FONT)
            rfonts.set(qn("w:hAnsi"), FONT)
            rfonts.set(qn("w:cs"), FONT)


def add_space_before_payment_milestones(doc: Document) -> None:
    """Spacing + hard page break before Payment milestones (matches HTML print/PDF)."""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t.startswith("Payment milestones"):
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.page_break_before = True
            break


def main() -> int:
    doc = Document(DOCX_PATH)
    style_doc_defaults(doc)

    for p in list(doc.paragraphs):
        if (p.text or "").strip() == "Print / Save as PDF":
            el = p._element
            el.getparent().remove(el)

    for p in doc.paragraphs:
        if p.style.name == "Heading 2":
            add_heading2_with_bar(p, BAR_IMAGE)

    apply_font_all_runs(doc, FONT)

    add_space_before_payment_milestones(doc)

    doc.save(DOCX_PATH)
    print("Updated:", DOCX_PATH)
    return 0


if __name__ == "__main__":
    sys.exit(main())
